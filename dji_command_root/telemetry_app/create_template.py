import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

def set_font(run, font_name='宋体', size=12):
    """强制设置中英文字体和大小，解决中文字体大小不一致的bug"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)

def create_template():
    # Ensure templates directory exists
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    templates_dir = os.path.join(base_dir, "templates")
    os.makedirs(templates_dir, exist_ok=True)

    template_path = os.path.join(templates_dir, "report_template.docx")

    doc = Document()

    # 修改正文默认样式，防止被覆盖
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    # Title
    title = doc.add_heading('巡检报告', 0)
    title.alignment = 1 # center

    # Basic Info
    doc.add_heading('一、任务基础信息', level=1)

    # 任务名称
    p1 = doc.add_paragraph()
    r1_label = p1.add_run('任务名称：')
    r1_label.bold = True
    r1_val = p1.add_run('{{ task_name }}')

    # 执行时间
    p2 = doc.add_paragraph()
    r2_label = p2.add_run('执行时间：')
    r2_label.bold = True
    r2_val = p2.add_run('{{ execute_time }}')

    # 任务类型
    p3 = doc.add_paragraph()
    r3_label = p3.add_run('任务类型：')
    r3_label.bold = True
    r3_val = p3.add_run('{{ task_type }}')

    # 统一强制设置字体和大小
    for r in [r1_label, r1_val, r2_label, r2_val, r3_label, r3_val]:
        set_font(r)

    # Statistics
    doc.add_heading('二、异常统计', level=1)
    p4 = doc.add_paragraph()
    r4_label = p4.add_run('总异常数：')
    r4_label.bold = True
    r4_val = p4.add_run('{{ total_anomalies }}')

    for r in [r4_label, r4_val]:
        set_font(r)

    # Anomalies List
    doc.add_heading('三、异常详情', level=1)

    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'

    # Row 0: Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '发现时间'
    hdr_cells[2].text = '任务类型'
    hdr_cells[3].text = '异常描述'
    hdr_cells[4].text = '位置(经纬度)'

    # Row 1: docxtpl loop 控制行 - {%tr %} 必须独占一行，docxtpl 会将整行替换为 Jinja2 标签
    for_loop_row = table.rows[1]
    for_loop_row.cells[0].text = '{%tr for item in anomalies %}'

    # Row 2: 数据行 - 包含实际的数据模板变量，会被 Jinja2 循环复制
    data_row = table.rows[2]
    data_row.cells[0].text = '{{ loop.index }}'
    data_row.cells[1].text = '{{ item.time }}'
    data_row.cells[2].text = '{{ item.level }}'
    data_row.cells[3].text = '{{ item.description }}'
    data_row.cells[4].text = '{{ item.longitude }}, {{ item.latitude }}'

    # Row 3: loop 结束控制行
    endfor_row = table.rows[3]
    endfor_row.cells[0].text = '{%tr endfor %}'

    # 设置表格的列宽，防止文字被挤成竖排
    table.columns[0].width = Cm(1.5) # 序号
    table.columns[1].width = Cm(4.0) # 发现时间
    table.columns[2].width = Cm(3.0) # 异常级别
    table.columns[3].width = Cm(5.0) # 异常描述
    table.columns[4].width = Cm(4.0) # 位置

    # 必须给每个 cell 单独设置宽度才生效
    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(4.0)
        row.cells[2].width = Cm(3.0)
        row.cells[3].width = Cm(5.0)
        row.cells[4].width = Cm(4.0)

    doc.save(template_path)
    print(f"Template created at {template_path}")

if __name__ == "__main__":
    create_template()
